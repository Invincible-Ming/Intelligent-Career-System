"""
文档上传、解析、切块和向量入库服务。
"""

from __future__ import annotations

import asyncio
import io
import mimetypes
import uuid
from pathlib import Path

import fitz
from docx import Document as WordDocument
from fastapi import UploadFile
from openpyxl import load_workbook
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.bm25_service import bm25_service
from app.bailian import bailian_service
from app.config import settings
from app.milvus_service import (
    DOCUMENT_TYPES,
    milvus_service,
)
from app.minio_service import minio_service
from app.models import Document

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".txt",
}


class DocumentService:
    """文档上传和知识库入库服务。"""

    async def upload_document(
            self,
            *,
            session: AsyncSession,
            file: UploadFile,
            document_type: str,
    ) -> Document:
        """上传文件并完成解析和向量入库。"""

        if document_type not in DOCUMENT_TYPES:
            raise ValueError(
                "document_type 必须是 "
                "resume、job_description 或 knowledge"
            )

        filename = file.filename or "document.txt"
        extension = Path(filename).suffix.lower()

        if extension not in ALLOWED_EXTENSIONS:
            raise ValueError(
                "仅支持 PDF、DOCX、XLSX 和 TXT 文件"
            )

        data = await file.read()

        if not data:
            raise ValueError("上传文件不能为空")

        if len(data) > settings.max_upload_size_bytes:
            raise ValueError(
                f"文件不能超过 {settings.MAX_UPLOAD_SIZE_MB} MB"
            )

        document_id = uuid.uuid4()
        content_type = (
                file.content_type
                or mimetypes.guess_type(filename)[0]
                or "application/octet-stream"
        )

        object_key = await minio_service.upload(
            document_id=str(document_id),
            filename=filename,
            data=data,
            content_type=content_type,
        )

        document = Document(
            id=document_id,
            filename=filename,
            document_type=document_type,
            minio_object_key=object_key,
            status="processing",
        )

        session.add(document)
        await session.commit()

        try:
            text = await asyncio.to_thread(
                parse_document,
                data,
                extension,
            )

            if not text.strip():
                raise ValueError(
                    "文档中没有提取到有效文本"
                )

            chunks = split_text(text)

            if not chunks:
                raise ValueError("文档切块结果为空")

            vectors = await embed_chunks(chunks)

            inserted_count = await milvus_service.insert(
                document_id=str(document_id),
                document_type=document_type,
                chunks=chunks,
                vectors=vectors,
            )

            await bm25_service.add_document(
                document_id=str(document_id),
                document_type=document_type,
                chunks=chunks,
            )

            document.status = "ready"
            document.chunk_count = inserted_count
            document.error_message = None

            await session.commit()
            await session.refresh(document)

            return document

        except Exception as exc:
            await session.rollback()

            failed_document = await session.get(
                Document,
                document_id,
            )

            if failed_document is not None:
                failed_document.status = "failed"
                failed_document.error_message = str(exc)[:1000]
                await session.commit()

            # 清理可能已经写入的部分向量。
            try:
                await milvus_service.delete_document(
                    str(document_id)
                )
                await bm25_service.delete_document(
                    str(document_id)
                )
            except Exception:
                pass

            raise RuntimeError(
                f"文档处理失败：{exc}"
            ) from exc

    async def get_document_text(
            self,
            *,
            session: AsyncSession,
            document_id: uuid.UUID,
    ) -> str:
        """从 MinIO 下载文档并重新解析文本。"""

        document = await session.get(
            Document,
            document_id,
        )

        if document is None:
            raise ValueError("文档不存在")

        if document.status != "ready":
            raise ValueError(
                f"文档当前不可用，状态：{document.status}"
            )

        data = await minio_service.download(
            document.minio_object_key
        )

        extension = Path(
            document.filename
        ).suffix.lower()

        text = await asyncio.to_thread(
            parse_document,
            data,
            extension,
        )

        if not text.strip():
            raise ValueError("文档中没有有效文本")

        return text

    async def list_documents(
            self,
            *,
            session: AsyncSession,
    ) -> list[Document]:
        """查询全部文档。"""

        result = await session.execute(
            select(Document).order_by(
                Document.created_at.desc()
            )
        )

        return list(result.scalars())

    async def delete_document(
            self,
            *,
            session: AsyncSession,
            document_id: uuid.UUID,
    ) -> None:
        """删除 PostgreSQL、Milvus 和 MinIO 中的文档。"""

        document = await session.get(
            Document,
            document_id,
        )

        if document is None:
            raise ValueError("文档不存在")

        await milvus_service.delete_document(
            str(document_id)
        )

        await bm25_service.delete_document(
            str(document_id)
        )

        await minio_service.delete(
            document.minio_object_key
        )

        await session.delete(document)
        await session.commit()


async def embed_chunks(
        chunks: list[str],
        batch_size: int = 10,
) -> list[list[float]]:
    """分批调用百炼 Embedding，避免单次输入过多。"""

    vectors: list[list[float]] = []

    for start in range(0, len(chunks), batch_size):
        batch = chunks[start: start + batch_size]
        batch_vectors = await bailian_service.embed(batch)
        vectors.extend(batch_vectors)

    return vectors


def parse_document(
        data: bytes,
        extension: str,
) -> str:
    """根据文件扩展名提取文本。"""

    if extension == ".pdf":
        return parse_pdf(data)

    if extension == ".docx":
        return parse_docx(data)

    if extension == ".xlsx":
        return parse_xlsx(data)

    if extension == ".txt":
        return parse_txt(data)

    raise ValueError(f"不支持的文件格式：{extension}")


def parse_pdf(data: bytes) -> str:
    """解析带文本层的 PDF。"""

    document = fitz.open(
        stream=data,
        filetype="pdf",
    )

    try:
        pages = [
            page.get_text("text").strip()
            for page in document
        ]
    finally:
        document.close()

    return "\n\n".join(
        page for page in pages if page
    )


def parse_docx(data: bytes) -> str:
    """解析 DOCX 段落和表格。"""

    document = WordDocument(io.BytesIO(data))

    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                parts.append(" | ".join(values))

    return "\n".join(parts)


def parse_xlsx(data: bytes) -> str:
    """解析 XLSX 中的非空单元格。"""

    workbook = load_workbook(
        io.BytesIO(data),
        read_only=True,
        data_only=True,
    )

    parts: list[str] = []

    try:
        for sheet in workbook.worksheets:
            parts.append(f"工作表：{sheet.title}")

            for row in sheet.iter_rows(
                    values_only=True
            ):
                values = [
                    str(value).strip()
                    for value in row
                    if value is not None
                ]

                if values:
                    parts.append(" | ".join(values))
    finally:
        workbook.close()

    return "\n".join(parts)


def parse_txt(data: bytes) -> str:
    """尝试使用常见中文文本编码读取 TXT。"""

    for encoding in (
            "utf-8-sig",
            "utf-8",
            "gb18030",
            "utf-16",
    ):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise ValueError("无法识别 TXT 文件编码")


def split_text(text: str) -> list[str]:
    """按照固定长度和重叠字符数切分文本。"""

    text = "\n".join(
        line.strip()
        for line in text.replace(
            "\r\n",
            "\n",
        ).splitlines()
        if line.strip()
    )

    if not text:
        return []

    chunks: list[str] = []
    start = 0

    while start < len(text):
        end = min(
            start + settings.CHUNK_SIZE,
            len(text),
        )

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = end - settings.CHUNK_OVERLAP

    return chunks


document_service = DocumentService()
